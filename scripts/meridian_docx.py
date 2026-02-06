#!/usr/bin/env python3
"""
Meridian (Nordic Slate) .docx generator.

Goals:
- Consistent typography defaults (fonts, sizes, colors) across the entire document.
- Elegant, professional headers/footers for Word -> PDF workflows.

Usage:
  uv sync
  uv run python scripts/meridian_docx.py --out out/meridian-tri-footer.docx --pdf
"""

from __future__ import annotations

import argparse
import datetime as dt
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

PALETTE = {
    "charcoal": "#2B2D32",
    "charcoal_light": "#3E4048",
    "green": "#3D6B5E",
    "green_light": "#4D8272",
    "stone": "#F4F3F0",
    "white": "#FFFFFF",
    "muted": "#6E7074",
    "border": "#D4D3D0",
}

FONTS = {
    "display": "Libre Baskerville",
    "body": "Karla",
    "mono": "IBM Plex Mono",
}


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


def _get_or_add(parent, tag: str):
    """
    Get-or-add an OOXML child element.

    `tag` must include the namespace prefix, e.g. "w:color".
    """
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_rfonts_all_scripts(rFonts, *, name: str) -> None:
    # Prefer explicit fonts over theme fonts for predictable Word/LibreOffice rendering.
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        attr = qn(theme_attr)
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]

    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def _set_style_run_defaults(
    style,
    *,
    font_name: str,
    size_pt: float | None = None,
    color_hex: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    style.font.name = font_name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if color_hex is not None:
        style.font.color.rgb = _rgb(color_hex)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if underline is not None:
        style.font.underline = underline

    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    _set_rfonts_all_scripts(rFonts, name=font_name)


def _set_run_font(
    run,
    *,
    name: str,
    size_pt: float | None = None,
    color_hex: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    font = run.font
    font.name = name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if color_hex is not None:
        font.color.rgb = _rgb(color_hex)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if underline is not None:
        font.underline = underline

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    _set_rfonts_all_scripts(rFonts, name=name)


def _apply_doc_defaults(
    doc: Document,
    *,
    font_name: str,
    size_pt: float,
    color_hex: str,
    line_spacing: float,
    space_after_pt: float,
) -> None:
    """
    Configure document-level defaults in styles.xml.

    This prevents stray theme fonts (e.g., Calibri) from appearing in fields, headers, etc.
    """
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    # Run defaults
    rPrDefault = _get_or_add(doc_defaults, "w:rPrDefault")
    rPr = _get_or_add(rPrDefault, "w:rPr")

    rFonts = _get_or_add(rPr, "w:rFonts")
    _set_rfonts_all_scripts(rFonts, name=font_name)

    # Word uses half-points (hps) in w:sz.
    hps = str(int(round(size_pt * 2)))
    sz = _get_or_add(rPr, "w:sz")
    sz.set(qn("w:val"), hps)
    szCs = _get_or_add(rPr, "w:szCs")
    szCs.set(qn("w:val"), hps)

    color = rPr.find(qn("w:color"))
    if color is None:
        # Place color near the other run defaults for readability.
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), color_hex.lstrip("#").upper())

    # Paragraph defaults
    pPrDefault = _get_or_add(doc_defaults, "w:pPrDefault")
    pPr = _get_or_add(pPrDefault, "w:pPr")
    spacing = _get_or_add(pPr, "w:spacing")

    # w:after uses twips (1/20 pt); w:line uses 240ths of a line when w:lineRule="auto".
    spacing.set(qn("w:after"), str(int(round(space_after_pt * 20))))
    spacing.set(qn("w:line"), str(int(round(line_spacing * 240))))
    spacing.set(qn("w:lineRule"), "auto")


def _add_field(paragraph, field_instr: str) -> None:
    """
    Insert a Word field (e.g., 'PAGE', 'NUMPAGES') into a paragraph.

    Word/LibreOffice renders the field result when opening/converting.
    """
    run = paragraph.add_run()
    r = run._r

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {field_instr} "

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_sep)
    r.append(fld_char_end)


def _set_paragraph_bottom_rule(paragraph, *, color_hex: str, size_eighth_pt: int = 4) -> None:
    """
    Add a subtle bottom rule to a paragraph (useful in headers).

    `size_eighth_pt` is in eighths of a point (w:sz). `4` ~= 0.5pt.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # Remove any existing paragraph border definition to avoid duplicates.
    for child in list(pPr):
        if child.tag == qn("w:pBdr"):
            pPr.remove(child)

    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighth_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.lstrip("#").upper())
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_document_setup(doc: Document, *, paper: str) -> None:
    section = doc.sections[0]

    # Margins: 1 inch / ~25mm all sides.
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Keep headers/footers comfortably inside the margin.
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    paper = paper.lower().strip()
    if paper == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    elif paper == "a4":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        raise ValueError(f"Unsupported paper: {paper!r} (expected 'letter' or 'a4')")


def apply_typography(doc: Document) -> None:
    # Defaults first: prevents theme fonts from leaking into fields and edge cases.
    _apply_doc_defaults(
        doc,
        font_name=FONTS["body"],
        size_pt=11,
        color_hex=PALETTE["charcoal"],
        line_spacing=1.5,
        space_after_pt=6,
    )

    styles = doc.styles

    # Normal body
    normal = styles["Normal"]
    _set_style_run_defaults(
        normal,
        font_name=FONTS["body"],
        size_pt=11,
        color_hex=PALETTE["charcoal"],
        bold=False,
    )
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    # Headings (Libre Baskerville headings are always regular-weight).
    h1 = styles["Heading 1"]
    _set_style_run_defaults(
        h1,
        font_name=FONTS["display"],
        size_pt=14,
        color_hex=PALETTE["charcoal"],
        bold=False,
    )
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.15

    h2 = styles["Heading 2"]
    _set_style_run_defaults(
        h2,
        font_name=FONTS["display"],
        size_pt=12,
        color_hex=PALETTE["charcoal"],
        bold=False,
    )
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15

    h3 = styles["Heading 3"]
    _set_style_run_defaults(
        h3,
        font_name=FONTS["body"],
        size_pt=11,
        color_hex=PALETTE["charcoal"],
        bold=True,  # Closest OOXML approximation for Karla SemiBold (600).
    )
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.2

    # Header/Footer built-in styles (used by Word UI too).
    for style_name, size_pt in (("Header", 8), ("Footer", 9)):
        try:
            st = styles[style_name]
        except KeyError:
            continue
        _set_style_run_defaults(
            st,
            font_name=FONTS["body"],
            size_pt=size_pt,
            color_hex=PALETTE["muted"],
            bold=False,
        )
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.0

    # Hyperlink: keep Word-friendly conventions but match brand color.
    try:
        hyperlink = styles["Hyperlink"]
        _set_style_run_defaults(
            hyperlink,
            font_name=FONTS["body"],
            color_hex=PALETTE["green"],
            underline=True,
        )
    except KeyError:
        pass


def apply_header_footer_tri_footer(
    doc: Document,
    *,
    brand_left: str,
    header_right: str,
    footer_left: str,
    footer_right: str,
    different_first_page: bool = True,
) -> None:
    """
    Preferred variant:
    - Header: left brand lockup + right doc title + subtle bottom rule.
    - Footer: left site + centered "Page X of Y" + right confidentiality label.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = different_first_page

    usable_width = section.page_width - section.left_margin - section.right_margin

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.style = "Header"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(4)
    hp.paragraph_format.tab_stops.clear_all()
    hp.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    r_left = hp.add_run(brand_left.upper())
    _set_run_font(r_left, name=FONTS["body"], bold=True)
    hp.add_run("\t")
    r_right = hp.add_run(header_right)
    _set_run_font(r_right, name=FONTS["body"])

    _set_paragraph_bottom_rule(hp, color_hex=PALETTE["border"], size_eighth_pt=4)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.style = "Footer"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.clear_all()
    fp.paragraph_format.tab_stops.add_tab_stop(int(usable_width / 2), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    fp.add_run(footer_left)
    fp.add_run("\t")
    fp.add_run("Page ")
    _add_field(fp, "PAGE")
    fp.add_run(" of ")
    _add_field(fp, "NUMPAGES")
    fp.add_run("\t")
    fp.add_run(footer_right)

    # Ensure fields adopt the footer font/family if Word/LO doesn't inherit it correctly.
    for r in fp.runs:
        _set_run_font(r, name=FONTS["body"])

    if different_first_page:
        # Common preference: no running header/footer on the title page.
        for p in section.first_page_header.paragraphs:
            p.text = ""
        for p in section.first_page_footer.paragraphs:
            p.text = ""


def add_title_page(
    doc: Document,
    *,
    overline: str,
    title: str,
    date_str: str,
    confidentiality: str,
) -> None:
    # Wordmark-ish text (we're not embedding actual logo assets here).
    p_wordmark = doc.add_paragraph()
    p_wordmark.paragraph_format.space_after = Pt(18)
    r = p_wordmark.add_run("MERIDIAN")
    _set_run_font(r, name=FONTS["display"], size_pt=24, color_hex=PALETTE["charcoal"], bold=False)

    p_overline = doc.add_paragraph()
    p_overline.paragraph_format.space_after = Pt(6)
    r = p_overline.add_run(overline.upper())
    _set_run_font(r, name=FONTS["body"], size_pt=9, color_hex=PALETTE["green"], bold=True)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(14)
    r = p_title.add_run(title)
    _set_run_font(r, name=FONTS["display"], size_pt=20, color_hex=PALETTE["charcoal"], bold=False)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(0)
    r = p_meta.add_run(f"{date_str}  ·  {confidentiality}")
    _set_run_font(r, name=FONTS["body"], size_pt=10, color_hex=PALETTE["muted"])

    doc.add_page_break()


def add_sample_content(doc: Document) -> None:
    doc.add_heading("From strategy to production", level=1)
    doc.add_paragraph(
        "Meridian builds healthcare AI systems with clinical credibility, technical depth, and quiet confidence."
    )

    doc.add_heading("Section heading", level=2)
    doc.add_paragraph(
        "This document is a starter template for Meridian-styled headers and footers generated via python-docx."
    )

    doc.add_heading("Small heading", level=3)
    doc.add_paragraph(
        "To force multiple pages for review, we include filler paragraphs below."
    )

    filler = "Clinicians will not tolerate workflow friction. If an AI system adds cognitive load, it fails."
    for _ in range(80):
        doc.add_paragraph(filler)


def convert_docx_to_pdf(docx_path: Path, *, out_pdf: Path | None = None) -> Path:
    soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    soffice_path = Path(soffice)
    if not soffice_path.exists():
        raise FileNotFoundError("LibreOffice 'soffice' not found in PATH or default macOS location.")

    docx_path = docx_path.resolve()
    if out_pdf is None:
        out_pdf = docx_path.with_suffix(".pdf")
    else:
        out_pdf = out_pdf.resolve()

    outdir = out_pdf.parent
    outdir.mkdir(parents=True, exist_ok=True)

    # LibreOffice writes <basename>.pdf into outdir.
    subprocess.run(
        [
            str(soffice_path),
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(outdir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )

    produced = outdir / f"{docx_path.stem}.pdf"
    if not produced.exists():
        raise FileNotFoundError("LibreOffice conversion succeeded but output PDF was not found.")

    if produced != out_pdf:
        produced.replace(out_pdf)

    return out_pdf


def build_doc(
    *,
    title: str,
    header_text: str,
    paper: str,
    overline: str,
    confidentiality: str,
    brand_left: str,
    footer_left: str,
    footer_right: str,
) -> Document:
    doc = Document()
    apply_document_setup(doc, paper=paper)
    apply_typography(doc)
    apply_header_footer_tri_footer(
        doc,
        brand_left=brand_left,
        header_right=header_text,
        footer_left=footer_left,
        footer_right=footer_right,
        different_first_page=True,
    )

    today = dt.date.today()
    # Format: "6 February 2026" (day month year, no commas).
    date_str = f"{today.day} {today.strftime('%B %Y')}"
    add_title_page(
        doc,
        overline=overline,
        title=title,
        date_str=date_str,
        confidentiality=confidentiality,
    )
    add_sample_content(doc)

    return doc


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--paper", choices=["letter", "a4"], default="letter")
    parser.add_argument("--title", default="Meridian Brand Guidelines")
    parser.add_argument("--header", default="", help="Running header text (defaults to --title)")
    parser.add_argument("--brand-left", default="Meridian", help="Header left lockup text")
    parser.add_argument("--footer-left", default="meridian.health", help="Footer left text")
    parser.add_argument("--footer-right", default="Confidential", help="Footer right text")
    parser.add_argument("--overline", default="Meridian")
    parser.add_argument(
        "--confidentiality",
        default="Confidential — For internal and approved partner use only",
    )
    parser.add_argument("--pdf", action="store_true", help="Also convert to PDF via LibreOffice")
    parser.add_argument("--out-pdf", type=Path, default=None, help="Optional output PDF path")
    args = parser.parse_args()

    out_docx = args.out
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    header_text = args.header or args.title
    doc = build_doc(
        title=args.title,
        header_text=header_text,
        paper=args.paper,
        overline=args.overline,
        confidentiality=args.confidentiality,
        brand_left=args.brand_left,
        footer_left=args.footer_left,
        footer_right=args.footer_right,
    )
    doc.save(out_docx)

    if args.pdf:
        convert_docx_to_pdf(out_docx, out_pdf=args.out_pdf)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
